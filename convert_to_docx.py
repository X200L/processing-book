#!/usr/bin/env python3
"""Конвертирует учебник по Processing из Markdown в Word (.docx)."""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("Ошибка: установите python-docx: pip install python-docx")
    sys.exit(1)


def setup_styles(doc):
    """Настраиваем базовые стили документа."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hname = f'Heading {level}'
        if hname in doc.styles:
            hstyle = doc.styles[hname]
            hstyle.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            hstyle.paragraph_format.space_before = Pt(12 if level > 1 else 18)
            hstyle.paragraph_format.space_after = Pt(6)


def add_code_block(doc, code_text, language=""):
    """Добавляет блок кода с серым фоном и monospace шрифтом."""
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        # серый фон
        shading = run._element.get_or_add_rPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'F0F0F0',
        })
        shading.append(shd)


def add_table_from_md(doc, md_table):
    """Парсит markdown-таблицу и добавляет её в документ."""
    lines = [l.strip() for l in md_table.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return

    header_line = lines[0]
    # separator line: lines[1]
    data_lines = lines[2:]

    headers = [h.strip().strip('|') for h in header_line.split('|') if h.strip()]
    rows = []
    for dl in data_lines:
        cells = [c.strip() for c in dl.split('|') if c.strip()]
        if cells:
            rows.append(cells)

    if not rows:
        return

    ncols = max(len(headers), max(len(r) for r in rows))
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def parse_markdown_to_docx(md_path, docx_path):
    """Читает .md и создаёт .docx."""
    doc = Document()
    setup_styles(doc)

    # Устанавливаем поля
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    text = Path(md_path).read_text(encoding='utf-8')

    # Удаляем [[MyArticles]] и оглавление (до первого ---)
    text = re.sub(r'\[\[MyArticles\]\]\s*', '', text)
    text = re.sub(r'## Оглавление\n.*?(?=---)', '', text, flags=re.DOTALL)

    # Разбиваем на блоки: код, таблицы, детали, остальное
    lines = text.split('\n')
    i = 0
    in_code = False
    code_buffer = []
    code_lang = ""
    in_details = False
    details_summary = ""
    details_buffer = []

    while i < len(lines):
        line = lines[i]

        # === Кодовые блоки ===
        if line.startswith('```'):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_buffer = []
                i += 1
                continue
            else:
                in_code = False
                add_code_block(doc, '\n'.join(code_buffer), code_lang)
                code_buffer = []
                i += 1
                continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # === details/summary ===
        if line.strip().startswith('<details>'):
            in_details = True
            details_buffer = []
            i += 1
            continue

        if line.strip().startswith('</details>'):
            in_details = False
            # Добавляем содержимое details как обычный текст
            for dl in details_buffer:
                if dl.strip().startswith('<summary>'):
                    m = re.search(r'<summary>(.*?)</summary>', dl)
                    if m:
                        p = doc.add_paragraph()
                        run = p.add_run(m.group(1))
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                elif dl.strip().startswith('- ') or dl.strip().startswith('* '):
                    p = doc.add_paragraph(dl.strip(), style='List Bullet')
                    for run in p.runs:
                        run.font.size = Pt(10)
                elif dl.strip().startswith('1.') or dl.strip().startswith('2.') or dl.strip().startswith('3.'):
                    p = doc.add_paragraph(dl.strip(), style='List Number')
                    for run in p.runs:
                        run.font.size = Pt(10)
                elif dl.strip():
                    p = doc.add_paragraph(dl.strip())
                    for run in p.runs:
                        run.font.size = Pt(10)
            i += 1
            continue

        if in_details:
            details_buffer.append(line)
            i += 1
            continue

        # === Таблицы ===
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s\|\-:]+$', lines[i + 1]):
            table_lines = []
            while i < len(lines) and ('|' in lines[i] or lines[i].strip() == ''):
                if lines[i].strip():
                    table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, '\n'.join(table_lines))
            continue

        # === Заголовки ===
        hm = re.match(r'^(#{1,3})\s+(.+)$', line)
        if hm:
            level = len(hm.group(1))
            title = hm.group(2).strip()
            p = doc.add_heading(title, level=level)
            i += 1
            continue

        # === Горизонтальная черта ===
        if re.match(r'^---\s*$', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(6)
            i += 1
            continue

        # === Маркированные списки ===
        if re.match(r'^[\s]*[-*]\s+', line):
            txt = re.sub(r'^[\s]*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_inline_formatting(p, txt)
            i += 1
            continue

        # === Нумерованные списки ===
        if re.match(r'^\s*\d+\.\s+', line):
            txt = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_inline_formatting(p, txt)
            i += 1
            continue

        # === Обычный параграф ===
        if line.strip():
            p = doc.add_paragraph()
            add_inline_formatting(p, line)
        else:
            doc.add_paragraph()  # пустая строка

        i += 1

    doc.save(docx_path)
    print(f"Готово: {docx_path}")


def add_inline_formatting(p, text):
    """Обрабатывает жирный текст (**), курсив (*), код (`) и ссылки."""
    # Разбиваем на части: **жирный**, *курсив*, `код`, [текст](url)
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))'
    last_end = 0
    for m in re.finditer(pattern, text):
        # текст до совпадения
        if m.start() > last_end:
            run = p.add_run(text[last_end:m.start()])
            run.font.size = Pt(11)

        if m.group(1) and m.group(2):  # **жирный**
            run = p.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(11)
        elif m.group(1) and m.group(3):  # *курсив*
            run = p.add_run(m.group(3))
            run.italic = True
            run.font.size = Pt(11)
        elif m.group(4):  # `код`
            run = p.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xE0, 0x40, 0x40)
        elif m.group(5) and m.group(6):  # [текст](url)
            run = p.add_run(m.group(5) + " (" + m.group(6) + ")")
            run.font.color.rgb = RGBColor(0x0E, 0x6E, 0xC0)
            run.font.size = Pt(11)
            run.underline = True

        last_end = m.end()

    if last_end < len(text):
        run = p.add_run(text[last_end:])
        run.font.size = Pt(11)


if __name__ == '__main__':
    md_file = '/Users/aleksandrvorobev/Documents/proc-lesson/book.md'
    docx_file = '/Users/aleksandrvorobev/Documents/proc-lesson/Учебник_Processing.docx'
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]

    parse_markdown_to_docx(md_file, docx_file)
